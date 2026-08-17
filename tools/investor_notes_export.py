"""
FIRE Capital Tools - investor update export.

PDF via matplotlib PdfPages and XLSX via openpyxl, the same two
mechanisms every other export in this app uses. Both libraries are
already dependencies.

WORD IS OFFERED, AND THE DEPENDENCY WAS APPROVED

Michelle asked for Word. python-docx was escalated rather than slipped
into a feature branch, and the escalation came back approved: MIT
licensed, pure Python, no system libraries. It is pinned in
requirements.txt.

Word is the format she actually forwards to investors, so .docx is a
first-class export here rather than a conversion of the PDF.

SECTIONS ARE CHOSEN PER UPDATE, NOT FIXED

Her real updates used a different set of sections each time, so nothing
here mandates a section list. build_docx() renders exactly the sections
it is handed, in the order it is handed them. select_sections() filters
an update's sections down to a caller's chosen keys without reordering.

There is deliberately no default subset. A caller that passes no
selection gets every section the update has -- "all" rather than a
curated few, because a hardcoded default is the fixed list this design
exists to avoid.

WHY THE PDF WRAPS TEXT BY HAND

PdfPages draws figures, not documents: matplotlib has no flow layout, so
a long paragraph does not reflow onto a second page by itself. The text
is wrapped and paginated here explicitly. That is the same approach
site_dd_report takes, and the reason both truncate visibly rather than
silently.
"""

from __future__ import annotations

import datetime
import textwrap
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt          # noqa: E402
from matplotlib.backends.backend_pdf import PdfPages   # noqa: E402

PAGE_SIZE = (8.5, 11)          # portrait: this is prose, not a dashboard
INK = "#1a2744"
MUTED = "#6b7280"
BODY = "#111827"

WRAP_AT = 96                   # characters per line at 9.5pt on portrait
LINES_PER_PAGE = 46
TOP = 0.90
LINE_STEP = 0.0185


def _lines_for(update: dict[str, Any], sections: list[dict[str, Any]],
               transcripts: list[dict[str, Any]]) -> list[tuple[str, str]]:
    """The whole document as (style, text) lines, ready to paginate."""
    out: list[tuple[str, str]] = []
    for section in sections:
        out.append(("h2", section["name"]))
        if section.get("empty"):
            out.append(("muted", section.get("empty_text") or "Not discussed."))
            out.append(("blank", ""))
            continue
        for point in section["points"]:
            wrapped = textwrap.wrap(point["text"], WRAP_AT - 2) or [""]
            out.append(("bullet", "• " + wrapped[0]))
            for cont in wrapped[1:]:
                out.append(("body", "  " + cont))
            out.append(("cite", f"    — {point.get('title')}, {point.get('date')}"))
        out.append(("blank", ""))

    out.append(("h2", "Sources"))
    for t in transcripts:
        out.append(("body",
                    f"  {t.get('transcript_date')}  "
                    f"{t.get('title') or t.get('original_name') or 'Untitled'}"))
    out.append(("blank", ""))
    out.append(("muted",
                "Every statement above is drawn from the meetings listed. "
                "Figures mentioned are as discussed on those calls and are not "
                "accounting records."))
    return out


def build_pdf(path, update: dict[str, Any], sections: list[dict[str, Any]],
              transcripts: list[dict[str, Any]]) -> Path:
    path = Path(path)
    label = update.get("property_label") or "Property"
    period = f"{update.get('period_start')} to {update.get('period_end')}"
    generated = (update.get("generated_at") or "")[:10]

    lines = _lines_for(update, sections, transcripts)
    pages = [lines[i:i + LINES_PER_PAGE]
             for i in range(0, len(lines), LINES_PER_PAGE)] or [[]]

    with PdfPages(str(path)) as pdf:
        for page_no, page in enumerate(pages, start=1):
            fig = plt.figure(figsize=PAGE_SIZE)
            fig.text(0.06, 0.955, "Investor Update", fontsize=16,
                     fontweight="bold", color=INK)
            fig.text(0.06, 0.930, label, fontsize=11, color="#4b5563")
            fig.text(0.94, 0.955, period, ha="right", fontsize=9.5, color=MUTED)
            fig.text(0.94, 0.930, f"Generated {generated}", ha="right",
                     fontsize=8.5, color=MUTED)

            y = TOP
            for style, text in page:
                if style == "blank":
                    y -= LINE_STEP
                    continue
                if style == "h2":
                    y -= LINE_STEP * 0.6
                    fig.text(0.06, y, text, fontsize=11.5, fontweight="bold",
                             color=INK)
                elif style == "cite":
                    fig.text(0.06, y, text, fontsize=7.8, color=MUTED,
                             style="italic")
                elif style == "muted":
                    fig.text(0.06, y, text, fontsize=8.5, color=MUTED)
                else:
                    fig.text(0.06, y, text, fontsize=9.5, color=BODY)
                y -= LINE_STEP

            fig.text(0.94, 0.04, f"Page {page_no} of {len(pages)}",
                     ha="right", fontsize=8, color=MUTED)
            pdf.savefig(fig)
            plt.close(fig)
    return path


def build_xlsx(path, update: dict[str, Any], sections: list[dict[str, Any]],
               transcripts: list[dict[str, Any]]) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font

    path = Path(path)
    wb = Workbook()
    ws = wb.active
    ws.title = "Investor Update"

    bold = Font(bold=True)
    wrap = Alignment(wrap_text=True, vertical="top")

    ws.append(["Investor Update"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([update.get("property_label") or ""])
    ws.append([f"{update.get('period_start')} to {update.get('period_end')}"])
    ws.append([f"Generated {(update.get('generated_at') or '')[:19]}"])
    ws.append([])

    ws.append(["Section", "Point", "Source meeting", "Date"])
    for cell in ws[ws.max_row]:
        cell.font = bold

    for section in sections:
        if section.get("empty"):
            ws.append([section["name"],
                       section.get("empty_text") or "Not discussed.", "", ""])
            continue
        for point in section["points"]:
            ws.append([section["name"], point["text"],
                       point.get("title"), point.get("date")])

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 90
    ws.column_dimensions["C"].width = 34
    ws.column_dimensions["D"].width = 12
    for row in ws.iter_rows(min_row=6, max_col=4):
        row[1].alignment = wrap

    src = wb.create_sheet("Sources")
    src.append(["Date", "Meeting", "Source", "File"])
    for cell in src[1]:
        cell.font = bold
    for t in transcripts:
        src.append([t.get("transcript_date"),
                    t.get("title") or "Untitled",
                    t.get("source"), t.get("original_name")])
    for col, width in zip("ABCD", (12, 40, 14, 40)):
        src.column_dimensions[col].width = width

    note = wb.create_sheet("Note")
    note["A1"] = ("Figures mentioned here are as discussed on the meetings "
                  "listed. They are narrative, not accounting records, and "
                  "nothing here has been written into Underwriting, Deal Dive "
                  "or any other tool.")
    note["A1"].alignment = wrap
    note.column_dimensions["A"].width = 100

    wb.save(str(path))
    return path



def select_sections(sections: list[dict[str, Any]],
                    keys: Any = None) -> list[dict[str, Any]]:
    """The sections a caller asked for, in the update's own order.

    `keys` None or empty means every section -- "all", not a default
    subset. Order comes from `sections` and never from `keys`, so a
    caller cannot reorder the document by reordering its request, and an
    unknown key is ignored rather than inventing an empty section.
    """
    if not keys:
        return list(sections)
    wanted = {str(k) for k in keys}
    return [s for s in sections if str(s.get("key")) in wanted]


def build_docx(path, update: dict[str, Any], sections: list[dict[str, Any]],
               transcripts: list[dict[str, Any]]) -> Path:
    """The same document the PDF renders, as a real Word file.

    Mirrors build_pdf()/build_xlsx(): same heading, same section order,
    same per-point attribution, same closing disclaimer. It is a separate
    renderer rather than a conversion because python-docx has flow layout
    and matplotlib does not -- this one gets real headings and real
    bullets that reflow when Michelle edits the file, which is the entire
    reason Word was asked for.

    An empty section is written with its "not discussed" line intact and
    visible. Nothing is generated to fill it.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    path = Path(path)
    label = update.get("property_label") or "Property"
    period = f"{update.get('period_start')} to {update.get('period_end')}"
    generated = (update.get("generated_at") or "")[:10]

    doc = Document()

    title = doc.add_heading("Investor Update", level=0)
    sub = doc.add_paragraph()
    sub.add_run(label).bold = True
    sub.add_run("\n" + period)
    meta = sub.add_run(f"\nGenerated {generated}")
    meta.font.size = Pt(8.5)
    meta.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    for section in sections:
        doc.add_heading(section["name"], level=1)
        if section.get("empty"):
            # Visibly empty. Never filled in -- a section Michelle enabled
            # and nobody discussed has to read as "nobody discussed it".
            run = doc.add_paragraph().add_run(
                section.get("empty_text") or "Not discussed.")
            run.italic = True
            run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
            continue
        for point in section["points"]:
            doc.add_paragraph(point["text"], style="List Bullet")
            cite = doc.add_paragraph()
            cite_run = cite.add_run(
                f"— {point.get('title')}, {point.get('date')}")
            cite_run.italic = True
            cite_run.font.size = Pt(8)
            cite_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_heading("Sources", level=1)
    for tr in transcripts:
        doc.add_paragraph(
            f"{tr.get('transcript_date')}  "
            f"{tr.get('title') or tr.get('original_name') or 'Untitled'}",
            style="List Bullet")

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Every statement above is drawn from the meetings listed. Figures "
        "mentioned are as discussed on those calls and are not accounting "
        "records.")
    note_run.font.size = Pt(8.5)
    note_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.save(str(path))
    return path


def suggested_filename(update: dict[str, Any], ext: str) -> str:
    label = "".join(c if c.isalnum() or c in " -_" else ""
                    for c in (update.get("property_label") or "property"))
    label = "-".join(label.split()).lower()[:48] or "property"
    return f"investor-update-{label}-{update.get('period_start')}.{ext}"
